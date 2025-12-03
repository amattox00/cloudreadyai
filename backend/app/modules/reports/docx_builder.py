from pathlib import Path
from typing import Any, Dict, List

from .utils import ensure_reports_dir

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
except ImportError:  # pragma: no cover
    Document = None  # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore


def _ensure_docx_available() -> None:
    if Document is None:
        raise RuntimeError(
            "python-docx is not installed. Install it with: pip install python-docx"
        )


def _add_title(document: "Document", title: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = run.font.size  # keep default size but bold
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if WD_ALIGN_PARAGRAPH else 1


def _add_section_heading(document: "Document", text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.upper())
    run.bold = True


def _add_paragraph(document: "Document", text: str) -> None:
    document.add_paragraph(text)


def _add_key_value_table(
    document: "Document", rows: List[Dict[str, Any]]
) -> None:
    """
    Very simple key-value style table.

    rows example:
        [{"label": "Total Servers", "value": 42}, ...]
    """
    if not rows:
        return

    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Light List" if "Light List" in document.styles else table.style

    for i, row in enumerate(rows):
        label_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]
        label_cell.text = str(row.get("label", ""))
        value_cell.text = str(row.get("value", ""))


def build_docx(
    run_id: str,
    report_type: str,
    context: Dict[str, Any],
    output_filename: str,
) -> Path:
    """
    Build a basic DOCX file for the given report_type and context.

    We intentionally keep this simple for Phase 1. We can refine styles
    and formatting in a later iteration.
    """
    _ensure_docx_available()

    reports_dir = ensure_reports_dir(run_id)
    output_path = reports_dir / output_filename

    document = Document()

    title = context.get("meta", {}).get("title", f"{report_type.title()} Report")
    client_name = context.get("meta", {}).get("client_name", "Client")

    _add_title(document, title)
    _add_paragraph(document, f"Client: {client_name}")
    _add_paragraph(document, "")

    # Executive overview
    exec_summary = context.get("executive_summary") or context.get(
        "summary", "Executive overview not available."
    )
    _add_section_heading(document, "Executive Overview")
    _add_paragraph(document, exec_summary)
    _add_paragraph(document, "")

    # High-level stats if present
    stats = context.get("environment", {}).get("stats") or []
    if isinstance(stats, list) and stats:
        _add_section_heading(document, "Environment Snapshot")
        _add_key_value_table(document, stats)
        _add_paragraph(document, "")

    # Recommendations if present
    recommendations = context.get("recommendations") or []
    if isinstance(recommendations, list) and recommendations:
        _add_section_heading(document, "Key Recommendations")
        for rec in recommendations:
            document.add_paragraph(f"- {rec}", style="List Bullet")

    document.save(str(output_path))
    return output_path
